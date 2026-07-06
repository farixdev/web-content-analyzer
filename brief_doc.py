# -*- coding: utf-8 -*-
"""
Generate a writer-ready SEO page-outline .docx from a search result's Competitor
Brief (search.search_and_analyze with deep analysis).

Profiles carry the brand rules (CTA, disclaimer, supported page types, middle-
section style). Two are built in - Intelinova (Australia, AI consulting) and
CanComCo (Canada, business communications) - plus a generic fallback for any
other brand name. The measurable tables (primary/secondary/entity densities,
competitor links) come from the real competitor research; the section headings
and FAQs are generated with Groq, grounded in the actual competitor terms. Body
copy is never written; a human writer fills every section.

Outline only. No emojis, no em dashes, no separator lines.
"""

import os
import re

import analyzer

# ----------------------------------------------------------------------------
# Brand profiles
# ----------------------------------------------------------------------------
_INSTR_MIDDLE = [
    "Should not be AI",
    "Should have 99 Grammarly Score",
    "Flesch reading ease Should have above 50 Readability Score",
    "Transition words must in 30% sentences (not 2 3 multiple)",
    "Passive voice use must be less the 10% sentences",
    "Consecutive sentences must not start with same words",
    "H1, H2, and H3 must have Primary keyword",
    "Grammarly Goals Selection: You need to select following Goals in the "
    "Grammarly while checking Grammarly Score",
    "Domain: Business",
    "Intent: Inform, Describe, Convince",
]

PROFILES = {
    "intelinova": {
        "brand": "Intelinova",
        "cta": "Talk To Expert",
        "market": "Australia",
        "audience_line": "Audience: Business Owners and Executives / Decision-Makers",
        "audience_prompt": "Australian business owners and executives / decision-makers",
        "page_types": ("service", "location", "industry"),
        "solutions_noun": "AI consulting",      # fixed disclaimer noun
        "use_service_noun": False,
        "middle_style": "services",             # Services (6), not Features
    },
    "cancomco": {
        "brand": "CanComCo",
        "cta": "Get A Quote",
        "market": "Canada",
        "audience_line": "Audience: Business Owners, Office Managers, and IT / Telecom Decision-Makers",
        "audience_prompt": "Canadian business owners, office managers, and IT/telecom decision-makers",
        "page_types": ("service", "location", "industry", "vendor", "comparison", "blog"),
        "solutions_noun": "business communications",
        "use_service_noun": True,               # disclaimer names the page's own service
        "middle_style": "features_addons",      # Features (10/6) or Add-ons (6)
    },
}

GENERIC_PROFILE = {
    "cta": "Get A Quote",
    "market": "",
    "audience_line": "Audience: Business Owners and Decision-Makers",
    "audience_prompt": "business owners and decision-makers",
    "page_types": ("service", "location", "industry", "vendor", "comparison", "blog"),
    "solutions_noun": "",
    "use_service_noun": True,
    "middle_style": "features_addons",
}


def get_profile(brand):
    """Resolve a brand name to a profile (known brand or generic)."""
    key = (brand or "").strip().lower().replace(" ", "")
    if key in PROFILES:
        return dict(PROFILES[key])
    p = dict(GENERIC_PROFILE)
    p["brand"] = (brand or "Brand").strip() or "Brand"
    p["solutions_noun"] = ""
    return p


# ----------------------------------------------------------------------------
# Page-type detection (fallback when Groq enrichment did not classify)
# ----------------------------------------------------------------------------
PLACES = {
    # Australia
    "sydney", "melbourne", "brisbane", "perth", "adelaide", "canberra", "hobart",
    "darwin", "gold coast", "australia", "nsw", "victoria", "queensland",
    # Canada
    "toronto", "montreal", "vancouver", "calgary", "edmonton", "ottawa", "quebec",
    "winnipeg", "hamilton", "mississauga", "ontario", "alberta", "canada",
    "british columbia", "manitoba", "near me",
}
INDUSTRIES = {
    "healthcare", "health", "medical", "dental", "retail", "ecommerce", "hospitality",
    "restaurant", "hotel", "legal", "law", "finance", "financial", "banking",
    "insurance", "manufacturing", "logistics", "construction", "education",
    "real estate", "property", "automotive", "energy", "government", "nonprofit",
    "mining", "agriculture", "pharma", "warehouse", "salon",
}
VENDORS = {
    "rogers", "bell", "telus", "shaw", "videotron", "ringcentral", "nextiva",
    "8x8", "ooma", "vonage", "dialpad", "gotoconnect", "zoom phone", "moneris",
    "clover", "square", "lightspeed", "toast", "shopify", "cisco", "yealink",
    "grandstream", "aws", "azure",
}


def detect_page_type(keyword):
    k = " " + (keyword or "").lower().strip() + " "
    if any(t in k for t in (" vs ", " versus ", " compared to ", " alternative", " alternatives")):
        return "comparison"
    if re.match(r"\s*(what|how|why|when|which|guide|checklist|is|are|do|does)\b", k) or k.strip().endswith("?"):
        return "blog"
    if any((" " + v + " ") in k for v in VENDORS):
        return "vendor"
    if any((" " + p + " ") in k or k.strip().endswith(p) for p in PLACES):
        return "location"
    if any((" " + i + " ") in k for i in INDUSTRIES):
        return "industry"
    return "service"


def _fmt_pct(x):
    try:
        x = round(float(x), 1)
    except Exception:
        return "0.5%"
    s = ("%.1f" % x).rstrip("0").rstrip(".")
    return (s or "0") + "%"


def _is_internet(product, primary):
    blob = ((product or "") + " " + (primary or "")).lower()
    return any(w in blob for w in ("internet", "broadband", "fibre", "fiber"))


def instructions_block(profile, service_noun):
    if profile.get("use_service_noun"):
        noun = service_noun or "business communications"
    else:
        noun = profile.get("solutions_noun") or "these"
    disclaimer = ("We don't deliver these solutions directly. %s is partner-led and "
                  "works with delivery partners who implement these %s solutions."
                  % (profile["brand"], noun))
    return [disclaimer] + _INSTR_MIDDLE + [profile["audience_line"]]


# ----------------------------------------------------------------------------
# Section generation (Groq, grounded in the real competitor terms; with fallback)
# ----------------------------------------------------------------------------
def _field_spec(page_type, middle_kind, middle_count):
    """[(key, count, description)] the outline needs for this page type."""
    if page_type in ("service", "location", "industry"):
        io = {"service": (6, "industries served or concrete buyer benefits"),
              "location": (4, "industries served or concrete buyer benefits"),
              "industry": (6, "specific use cases for this industry")}[page_type]
        mid_desc = {"features": "concrete factual product features",
                    "addons": "optional factual add-ons for this product",
                    "services": "service categories around the primary keyword"}[middle_kind]
        return [
            ("pains", 6, "problems the buyer feels without this, framed empathetically"),
            ("middle", middle_count, mid_desc),
            ("industries_or_usecases", io[0], io[1]),
            ("usps", 6, "reasons to choose the brand"),
            ("faqs", 5 if page_type == "location" else 6, "real natural buyer questions"),
        ]
    if page_type == "vendor":
        return [
            ("plans_features", 6, "real plan tiers or features of the vendor product, factual"),
            ("integrations", 6, "real verified integrations of the vendor product, factual"),
            ("why_through", 6, "reasons to buy this vendor product through the brand"),
            ("industries", 4, "buyer types or industries the product suits"),
            ("vs_direct", 6, "balanced points comparing buying through the brand vs buying direct"),
            ("usps", 6, "reasons to choose the brand"),
            ("faqs", 6, "real natural buyer questions"),
        ]
    if page_type == "comparison":
        return [
            ("criteria", 6, "dimensions to compare the two options on"),
            ("where_each_wins", 6, "strengths, three for each option"),
            ("fit", 4, "which type of buyer each option fits"),
            ("usps", 6, "reasons to choose the brand"),
            ("faqs", 6, "real natural buyer questions"),
        ]
    # blog
    return [
        ("takeaways", 5, "key takeaways, 3 to 7 words each"),
        ("body_sections", 7, "H2 body-section headings in reading order; the first must contain the primary keyword"),
        ("tips", 6, "practical checklist tips"),
        ("mistakes", 5, "common mistakes to avoid"),
        ("usps", 4, "soft non-salesy ways the brand helps"),
        ("faqs", 6, "real natural buyer questions"),
    ]


def _fallback_sections(spec, primary, secondary, entities):
    pool = [s["phrase"].title() for s in secondary] + [e["entity"].title() for e in entities]
    pool = [p for p in pool if p] or [primary.title()]
    out = {}
    for key, count, _ in spec:
        if key == "faqs":
            out[key] = ([
                "What is %s?" % primary,
                "How much does %s cost?" % primary,
                "How long does setup take?",
                "Do you support businesses like mine?",
                "What makes your approach different?",
                "How do we get started?",
            ] * 2)[:count]
        else:
            out[key] = [(pool[i % len(pool)])[:60] for i in range(count)]
    return out


def generate_sections(profile, page_type, primary, product, middle_kind, middle_count,
                      secondary, entities, api_key="", model=analyzer.DEFAULT_MODEL, progress=None):
    log = progress if callable(progress) else (lambda m: None)
    spec = _field_spec(page_type, middle_kind, middle_count)

    if not (api_key or "").strip():
        log("No Groq key: using derived headings.")
        return _fallback_sections(spec, primary, secondary, entities)

    sec_list = ", ".join(s["phrase"] for s in secondary[:12]) or "(none)"
    ent_list = ", ".join(e["entity"] for e in entities[:16]) or "(none)"
    keys_desc = "\n".join(
        "  %s (array of %d %s: %s)" % (k, n, "questions" if k == "faqs" else "strings 3-5 words each", d)
        for k, n, d in spec)

    system = (
        "You are an expert SEO content strategist producing a page OUTLINE only, "
        "never body copy. Reply with one valid JSON object, no prose, no markdown, "
        "no code fences. No emojis, no em dashes. Every heading is 3 to 5 words "
        "(FAQs are natural questions). Be specific and factual, grounded in the "
        "real competitor terms provided. Audience: %s." % profile["audience_prompt"]
    )
    user = (
        "Brand: %s (partner-led). Market: %s.\n"
        "Primary keyword: %s\n"
        "Product/service: %s\n"
        "Page type: %s\n"
        "Real secondary keywords from competitors: %s\n"
        "Real entities from competitors: %s\n\n"
        "Return a JSON object with EXACTLY these keys:\n%s\n"
        "Do not stuff the primary keyword into every FAQ."
    ) % (profile["brand"], profile.get("market") or "n/a", primary, product or primary,
         page_type, sec_list, ent_list, keys_desc)

    log("Generating outline headings with Groq (%s)..." % model)
    try:
        data = analyzer.groq_json(api_key.strip(), model, system, user, temperature=0.5, max_tokens=2000)
    except Exception as e:
        log("Groq generation failed (%s); using derived headings." % e)
        return _fallback_sections(spec, primary, secondary, entities)

    fb = _fallback_sections(spec, primary, secondary, entities)
    out = {}
    for key, count, _ in spec:
        vals = data.get(key) if isinstance(data, dict) else None
        vals = [str(v).strip() for v in (vals or []) if str(v).strip()]
        if not vals:
            vals = fb[key]
        out[key] = (vals + vals)[:count]
    return out


# ----------------------------------------------------------------------------
# Doc builder (house style)
# ----------------------------------------------------------------------------
def _slug(s):
    return re.sub(r"[^a-z0-9]+", "-", (s or "").lower()).strip("-") or "page"


def build_docx(path, profile, page_type, primary_keyword, service_noun, location,
               primary_density, secondary_rows, entity_rows, competitors, sections,
               middle_kind):
    from docx import Document
    from docx.shared import Pt, Inches

    brand = profile["brand"]
    cta_text = profile["cta"]
    pk = primary_keyword
    doc = Document()
    base = doc.styles["Normal"]; base.font.name = "Arial"; base.font.size = Pt(11)

    def header(text):
        p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.name = "Arial"

    def line(text):
        p = doc.add_paragraph(); r = p.add_run(text); r.font.name = "Arial"; r.font.size = Pt(11)

    def bullets(items):
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    def table2(h1, h2, rows):
        t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
        c = t.rows[0].cells; c[0].text = h1; c[1].text = h2
        for cc in c:
            for p in cc.paragraphs:
                for r in p.runs:
                    r.bold = True; r.font.name = "Arial"
        for a, b in rows:
            rc = t.add_row().cells; rc[0].text = str(a); rc[1].text = str(b)
            for cc in rc:
                for p in cc.paragraphs:
                    for r in p.runs:
                        r.font.name = "Arial"; r.font.size = Pt(10)
        for row in t.rows:
            row.cells[0].width = Inches(3.6); row.cells[1].width = Inches(2.2)
        doc.add_paragraph()

    def cta():
        header("CTA Section")
        line("Headline: (writer)"); line("Subheadline: (writer)"); line("CTA: %s" % cta_text)
        doc.add_paragraph()

    # ---- Shared head ----
    header("Instructions (place at top of page)")
    for li in instructions_block(profile, service_noun):
        line(li)
    doc.add_paragraph()

    header("Primary Keyword and Density")
    table2("Primary Keyword", "Density", [(pk, primary_density)])
    header("Secondary Keywords and Density")
    table2("Keywords", "Density", secondary_rows or [("(none captured)", "")])
    header("Entities and Density")
    table2("Entities", "Density", entity_rows or [("(none captured)", "")])

    header("Competitor Links (from the live SERP)")
    bullets(competitors or ["(none)"]); doc.add_paragraph()

    header("Hero Section")
    line('Headline: (writer, must contain "%s")' % pk)
    line("Subheadline: (writer, 30 words)")
    line("CTA: %s" % cta_text)
    doc.add_paragraph()

    # ---- Type-specific body ----
    if page_type in ("service", "location", "industry"):
        header("Why Businesses Need %s" % pk)
        line("Headline: (writer, must contain primary keyword)")
        line("Six pain-point bullets (writer writes 30 words under each):")
        bullets(sections["pains"]); doc.add_paragraph()

        if middle_kind == "addons":
            header("%s Add-ons" % pk)
            line("Headline: (writer, must contain primary keyword)")
            line("Six add-on bullets (writer writes 30 words under each). Factual only:")
        elif middle_kind == "services":
            header("%s Services" % pk)
            line("Headline: (writer, must contain primary keyword)")
            line("Six service bullets (writer writes 40 words under each):")
        elif page_type == "industry":
            header("Most Relevant %s Features" % pk)
            line("Headline: (writer, must contain primary keyword)")
            line("Most-relevant feature bullets matched to this industry (writer writes 40 words under each):")
        else:
            header("Top %s Features" % pk)
            line("Headline: (writer, must contain primary keyword)")
            line("Top ten feature bullets (writer writes 40 words under each):")
        bullets(sections["middle"]); doc.add_paragraph()
        cta()

        mid = sections["industries_or_usecases"]
        if page_type == "industry":
            header("%s Use Cases" % pk)
            line("Six use-case bullets (writer writes 30 words under each):")
        elif page_type == "location":
            header("Industries We Serve With %s" % pk)
            line("Four benefit bullets (writer writes 30 words under each):")
        else:
            header("Industries We Serve With %s" % pk)
            line("Six benefit bullets (writer writes 30 words under each):")
        bullets(mid); doc.add_paragraph()

        header("Why Choose %s for %s" % (brand, pk))
        line("Six USP bullets (writer writes 30 words under each):")
        bullets(sections["usps"]); doc.add_paragraph()
        cta()

    elif page_type == "vendor":
        header("%s Overview" % pk)
        line("Headline: (writer, must contain primary keyword)")
        line("Writer writes a 40 word plain overview of the vendor product. No bullets here.")
        doc.add_paragraph()
        header("%s Plans And Features" % pk)
        line("Headline: (writer, must contain primary keyword)")
        line("Six plan/feature bullets (writer writes 30 words under each). Factual only:")
        bullets(sections["plans_features"]); doc.add_paragraph()
        header("%s Integrations" % pk)
        line("Six integration bullets (writer writes 30 words under each). Factual only, real integrations:")
        bullets(sections["integrations"]); doc.add_paragraph()
        header("Why Get %s Through %s" % (pk, brand))
        line("Headline: (writer, must contain primary keyword)")
        line("Six partner-value bullets (writer writes 30 words under each):")
        bullets(sections["why_through"]); doc.add_paragraph()
        cta()
        header("Who It Is For / Industries Served")
        line("Four buyer/industry-fit bullets (writer writes 30 words under each):")
        bullets(sections["industries"]); doc.add_paragraph()
        header("%s vs Buying Direct" % brand)
        line("Headline: (writer, must contain primary keyword)")
        line("Six balanced comparison bullets (writer writes 30 words under each):")
        bullets(sections["vs_direct"]); doc.add_paragraph()
        header("Why Choose %s for %s" % (brand, pk))
        line("Six USP bullets (writer writes 30 words under each):")
        bullets(sections["usps"]); doc.add_paragraph()
        cta()

    elif page_type == "comparison":
        header("Quick Verdict (%s)" % pk)
        line("Headline: (writer, must contain primary keyword)")
        line("Writer writes a 30 word at-a-glance summary. No bullets here.")
        doc.add_paragraph()
        header("Comparison Criteria (%s)" % pk)
        line("Headline: (writer, must contain primary keyword)")
        line("Six comparison-dimension bullets (writer writes 30 words under each):")
        bullets(sections["criteria"]); doc.add_paragraph()
        header("Where Each Option Wins (%s)" % pk)
        line("Headline: (writer, must contain primary keyword)")
        line("Six bullets, three strengths per option (writer writes 30 words under each):")
        bullets(sections["where_each_wins"]); doc.add_paragraph()
        cta()
        header("Which Fits Your Business (%s)" % pk)
        line("Four buyer-fit bullets (writer writes 30 words under each):")
        bullets(sections["fit"]); doc.add_paragraph()
        header("Why Choose %s for %s" % (brand, pk))
        line("Six USP bullets (writer writes 30 words under each):")
        bullets(sections["usps"]); doc.add_paragraph()
        cta()

    elif page_type == "blog":
        header("Key Takeaways")
        line("Each takeaway is a 3 to 7 word heading (writer writes one sentence under each):")
        bullets(sections["takeaways"]); doc.add_paragraph()
        header("Article Body Sections")
        line("First H2 must contain the primary keyword; the rest read naturally.")
        line("H2 body-section headings in reading order (writer writes each section):")
        bullets(sections["body_sections"]); doc.add_paragraph()
        header("Practical Tips / Checklist")
        line("Actionable checklist bullets (writer writes one sentence under each):")
        bullets(sections["tips"]); doc.add_paragraph()
        header("Common Mistakes To Avoid")
        line("Mistake bullets (writer writes one sentence under each):")
        bullets(sections["mistakes"]); doc.add_paragraph()
        header("How %s Helps" % brand)
        line("Four soft, non-salesy USP bullets (writer writes 30 words under each):")
        bullets(sections["usps"]); doc.add_paragraph()
        cta()

    # ---- FAQs (all types). Location pages get a required area FAQ first. ----
    header("FAQs: %s" % pk)
    faqs = list(sections["faqs"])
    if page_type == "location":
        loc = location or "this area"
        faqs = [("Which areas do you serve in %s? (REQUIRED: writer's answer must "
                 "list every neighbourhood and area served)" % loc)] + faqs
        line("First FAQ is required (areas served, answer lists all neighbourhoods); writer writes answers:")
    else:
        line("Six real buyer questions (writer writes answers):")
    bullets(faqs)

    doc.save(path)
    return path


# ----------------------------------------------------------------------------
# Orchestrator
# ----------------------------------------------------------------------------
LABEL = {"service": "Service", "location": "Location-Service", "industry": "Industry",
         "vendor": "Vendor", "comparison": "Comparison", "blog": "Blog"}


def resolve_page_type(search_result, brand, page_type=None):
    """Final page type: explicit override, else Groq enrichment, else heuristic,
    clamped to what the brand profile supports."""
    profile = get_profile(brand)
    agg = search_result.get("aggregate") or {}
    pt = (page_type or agg.get("page_type") or detect_page_type(search_result["query"]) or "service").lower()
    if pt not in profile["page_types"]:
        # Fall back to the closest supported type.
        for alt in ("service", "industry", "location"):
            if alt in profile["page_types"]:
                pt = alt
                break
    return pt


def default_filename(search_result, brand, page_type):
    return "%s-%s-%s.docx" % (get_profile(brand)["brand"].replace(" ", ""),
                              LABEL.get(page_type, "Page"), _slug(search_result["query"]))


def generate_brief_docx(search_result, path, brand="Intelinova", api_key="",
                        model=analyzer.DEFAULT_MODEL, page_type=None, progress=None):
    """Build the outline .docx from a search result that has an aggregate block."""
    log = progress if callable(progress) else (lambda m: None)
    agg = search_result.get("aggregate")
    if not agg:
        raise RuntimeError(
            "No competitor brief to build from. Set 'Deep-analyze top N' to 3-5 "
            "and search again, then export the brief.")

    profile = get_profile(brand)
    primary = search_result["query"]
    page_type = resolve_page_type(search_result, brand, page_type)
    product = agg.get("product", "")
    service_noun = agg.get("service_noun") or product or primary
    location = agg.get("location", "")
    log("Brand: %s | page type: %s | product: %s" % (profile["brand"], page_type, product or "n/a"))

    # Middle-section kind + count.
    if profile["middle_style"] == "services":
        middle_kind, middle_count = "services", 6
    elif _is_internet(product, primary):
        middle_kind, middle_count = "addons", 6
    elif page_type == "industry":
        middle_kind, middle_count = "features", 6
    else:
        middle_kind, middle_count = "features", 10

    t = agg["targets"]
    primary_density = _fmt_pct(t["primary"]["target_density_pct"])
    secondary_rows = [(s["phrase"], _fmt_pct(s["target_density_pct"])) for s in t["secondary"]]
    entity_rows = [(e["entity"], _fmt_pct(e["target_density_pct"])) for e in t["entities"]]
    competitors = [r["url"] for r in search_result["results"] if r.get("analysis")][:5]
    if not competitors:
        competitors = [r["url"] for r in search_result["results"][:5]]

    sections = generate_sections(
        profile, page_type, primary, product, middle_kind, middle_count,
        agg["secondary_keywords"], agg["entities"], api_key=api_key, model=model, progress=progress)

    log("Writing document...")
    return build_docx(path, profile, page_type, primary, service_noun, location,
                      primary_density, secondary_rows, entity_rows, competitors, sections,
                      middle_kind)
